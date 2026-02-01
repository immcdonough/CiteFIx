"""Document parsing service for extracting text and citations from .docx files."""

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from docx import Document
from docx.text.paragraph import Paragraph


@dataclass
class DocumentSection:
    """Represents a section of the document."""
    title: str
    content: str
    start_pos: int
    end_pos: int
    paragraphs: list[str]


@dataclass
class ParsedDocument:
    """Parsed document with extracted sections."""
    full_text: str
    paragraphs: list[str]
    references_section: Optional[DocumentSection]
    reference_entries: list[str]
    body_text: str  # Text excluding references


# Common reference section headers
REFERENCE_HEADERS = [
    r"references?",
    r"bibliography",
    r"works?\s+cited",
    r"literature\s+cited",
    r"sources?",
    r"citations?",
]

REFERENCE_PATTERN = re.compile(
    r"^\s*(" + "|".join(REFERENCE_HEADERS) + r")\s*$",
    re.IGNORECASE
)


def parse_docx(file_path: Path | str) -> ParsedDocument:
    """
    Parse a .docx file and extract text and reference section.

    Args:
        file_path: Path to the .docx file

    Returns:
        ParsedDocument with extracted content
    """
    doc = Document(file_path)

    paragraphs: list[str] = []
    full_text_parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
            full_text_parts.append(text)

    full_text = "\n\n".join(full_text_parts)

    # Find references section
    references_section = None
    reference_entries: list[str] = []
    body_paragraphs: list[str] = []
    ref_start_idx = None

    for idx, para_text in enumerate(paragraphs):
        if REFERENCE_PATTERN.match(para_text):
            ref_start_idx = idx
            break
        body_paragraphs.append(para_text)

    if ref_start_idx is not None:
        # Extract reference entries (everything after the header)
        ref_paragraphs = paragraphs[ref_start_idx + 1:]
        reference_entries = _parse_reference_entries(ref_paragraphs)

        ref_content = "\n".join(ref_paragraphs)
        body_text = "\n\n".join(body_paragraphs)

        # Calculate positions
        body_len = len(body_text)
        references_section = DocumentSection(
            title=paragraphs[ref_start_idx],
            content=ref_content,
            start_pos=body_len,
            end_pos=len(full_text),
            paragraphs=ref_paragraphs,
        )
    else:
        body_text = full_text

    return ParsedDocument(
        full_text=full_text,
        paragraphs=paragraphs,
        references_section=references_section,
        reference_entries=reference_entries,
        body_text=body_text,
    )


def _parse_reference_entries(paragraphs: list[str]) -> list[str]:
    """
    Parse individual reference entries from reference section paragraphs.

    Handles common formats:
    - Each paragraph is one reference
    - Numbered references: 1. Author...
    - Bracketed: [1] Author...
    """
    entries: list[str] = []
    current_entry: list[str] = []

    # Pattern for new reference start
    # Handles hyphenated names like "Fernandez-Mendoza J" and various Unicode dashes
    # Also handles names with apostrophes like "O'Connor J"
    new_ref_pattern = re.compile(
        r"^(\[\d+\]|\d+\.\s+|[A-Z][a-zA-Z'\-\u2010\u2011\u2012\u2013\u2014]+,?\s+[A-Z])"
    )

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Check if this starts a new reference
        if new_ref_pattern.match(para):
            if current_entry:
                entries.append(" ".join(current_entry))
            current_entry = [para]
        else:
            # Continuation of previous reference
            current_entry.append(para)

    # Don't forget the last entry
    if current_entry:
        entries.append(" ".join(current_entry))

    return entries


def update_docx_references(
    input_path: Path | str,
    output_path: Path | str,
    updated_references: list[str],
) -> None:
    """
    Update the references section in a .docx file.

    Args:
        input_path: Path to original .docx
        output_path: Path for output .docx
        updated_references: New reference entries
    """
    doc = Document(input_path)

    # Find references section
    ref_start_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if REFERENCE_PATTERN.match(para.text.strip()):
            ref_start_idx = idx
            break

    if ref_start_idx is None:
        # No references section found - append one
        doc.add_paragraph("References")
        for ref in updated_references:
            doc.add_paragraph(ref)
    else:
        # Clear existing references and add new ones
        # Keep the header paragraph
        ref_paragraphs = doc.paragraphs[ref_start_idx + 1:]

        # Remove old reference paragraphs
        for para in ref_paragraphs:
            p = para._element
            p.getparent().remove(p)

        # Add new references after the header
        # We need to insert after the references header
        ref_header = doc.paragraphs[ref_start_idx]

        # Add new reference paragraphs
        for ref in updated_references:
            new_para = doc.add_paragraph(ref)

    doc.save(output_path)


def extract_text_with_positions(doc: Document) -> list[tuple[str, int, int]]:
    """
    Extract text from document with character positions.

    Returns:
        List of (text, start_pos, end_pos) tuples for each paragraph
    """
    result: list[tuple[str, int, int]] = []
    current_pos = 0

    for para in doc.paragraphs:
        text = para.text
        if text.strip():
            start = current_pos
            end = current_pos + len(text)
            result.append((text, start, end))
            current_pos = end + 2  # Account for paragraph break

    return result
